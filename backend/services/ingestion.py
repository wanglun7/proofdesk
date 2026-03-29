import json
import os
import tempfile
from pathlib import Path

import dashscope
from dashscope import TextEmbedding
from sqlalchemy.ext.asyncio import AsyncSession

from models import Document, Chunk
from config import settings


# --- Text extraction ---

def extract_text_from_pdf(path: str) -> list[tuple[int, str]]:
    """Returns list of (page_num, text)."""
    from llama_index.readers.file import PDFReader
    docs = PDFReader().load_data(Path(path))
    return [(i, d.text) for i, d in enumerate(docs)]


def extract_text_from_docx(path: str) -> list[tuple[int, str]]:
    from docx import Document as DocxDoc
    doc = DocxDoc(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if not paragraphs:
        return [(0, '')]
    # Group into chunks of 30 paragraphs so we get meaningful chunking progress
    GROUP = 30
    pages = []
    for i in range(0, len(paragraphs), GROUP):
        group_text = "\n".join(paragraphs[i:i + GROUP])
        pages.append((i // GROUP, group_text))
    return pages


def extract_text_from_txt(path: str) -> list[tuple[int, str]]:
    text = Path(path).read_text(encoding="utf-8")
    lines = [l for l in text.splitlines() if l.strip()]
    if not lines:
        return [(0, '')]
    GROUP = 50
    pages = []
    for i in range(0, len(lines), GROUP):
        pages.append((i // GROUP, "\n".join(lines[i:i + GROUP])))
    return pages


def extract_text_from_pptx(path: str) -> list[tuple[int, str]]:
    from llama_index.readers.file import PptxReader
    docs = PptxReader().load_data(Path(path))
    return [(i, d.text) for i, d in enumerate(docs)]


def extract_text_from_html(path: str) -> list[tuple[int, str]]:
    from llama_index.readers.file import HTMLTagReader
    docs = HTMLTagReader().load_data(Path(path))
    return [(0, "\n".join(d.text for d in docs))]


def extract_text_from_markdown(path: str) -> list[tuple[int, str]]:
    from llama_index.readers.file import MarkdownReader
    docs = MarkdownReader().load_data(Path(path))
    return [(0, "\n".join(d.text for d in docs))]


def extract_text(path: str, file_type: str) -> list[tuple[int, str]]:
    if file_type == "pdf":
        return extract_text_from_pdf(path)
    elif file_type in ("docx", "doc"):
        return extract_text_from_docx(path)
    elif file_type == "pptx":
        return extract_text_from_pptx(path)
    elif file_type in ("html", "htm"):
        return extract_text_from_html(path)
    elif file_type in ("md", "markdown"):
        return extract_text_from_markdown(path)
    else:
        return extract_text_from_txt(path)


# --- Sentence-window chunking ---

def sentence_window_chunks(text: str, window_size: int = 3) -> list[tuple[str, str]]:
    """
    Returns list of (sentence, window) pairs using LlamaIndex SentenceWindowNodeParser.
    - sentence: the individual sentence (used for embedding — precise retrieval)
    - window: surrounding context sentences (sent to LLM for generation)
    """
    from llama_index.core.node_parser import SentenceWindowNodeParser
    from llama_index.core.schema import Document as LIDoc

    parser = SentenceWindowNodeParser.from_defaults(
        window_size=window_size,
        window_metadata_key="window",
        original_text_metadata_key="original_sentence",
    )
    nodes = parser.get_nodes_from_documents([LIDoc(text=text)])
    result = []
    for node in nodes:
        sentence = node.text.strip()
        window = node.metadata.get("window", sentence).strip()
        if sentence:
            result.append((sentence, window))
    return result


# --- Embedding ---

def embed_batch(texts: list[str]) -> list[list[float]]:
    """Embed a single batch (≤10 texts)."""
    dashscope.api_key = settings.dashscope_api_key
    resp = TextEmbedding.call(
        model=settings.embedding_model,
        input=texts,
        dimension=settings.embed_dim,
    )
    if resp.status_code != 200:
        raise RuntimeError(f"Embedding API error: {resp.message}")
    return [e["embedding"] for e in resp.output["embeddings"]]


def embed_texts(texts: list[str]) -> list[list[float]]:
    all_embeddings = []
    for i in range(0, len(texts), 10):
        all_embeddings.extend(embed_batch(texts[i:i + 10]))
    return all_embeddings


# --- Streaming ingestion pipeline ---

async def ingest_file_stream(
    file_bytes: bytes, filename: str, db: AsyncSession, project_id: str | None = None
):
    """Async generator yielding SSE-formatted progress strings."""
    import asyncio
    BATCH = 10
    suffix = Path(filename).suffix.lower().lstrip(".")
    loop = asyncio.get_event_loop()

    with tempfile.NamedTemporaryFile(suffix=f".{suffix}", delete=False) as f:
        f.write(file_bytes)
        tmp_path = f.name

    try:
        yield f"data: {json.dumps({'type': 'extracting'})}\n\n"
        # Run blocking extraction in thread pool so event loop stays free
        pages = await loop.run_in_executor(None, extract_text, tmp_path, suffix)

        yield f"data: {json.dumps({'type': 'chunking', 'current': 0, 'total': len(pages)})}\n\n"
        all_sentences: list[str] = []
        all_windows: list[str] = []
        all_pages: list[int] = []
        for j, (page_num, page_text) in enumerate(pages):
            # Run blocking chunking in thread pool
            pairs = await loop.run_in_executor(None, sentence_window_chunks, page_text)
            for sentence, window in pairs:
                all_sentences.append(sentence)
                all_windows.append(window)
                all_pages.append(page_num)
            yield f"data: {json.dumps({'type': 'chunking', 'current': j + 1, 'total': len(pages)})}\n\n"

        total = len(all_sentences)
        yield f"data: {json.dumps({'type': 'chunking_done', 'total': total})}\n\n"

        if total == 0:
            yield f"data: {json.dumps({'type': 'error', 'error': 'No text content extracted from file'})}\n\n"
            return

        # Create DB record
        doc = Document(filename=filename, file_type=suffix, project_id=project_id)
        db.add(doc)
        await db.flush()

        # Embed batch-by-batch in thread pool, yield progress after each batch
        embeddings: list[list[float]] = []
        for i in range(0, total, BATCH):
            batch = all_sentences[i:i + BATCH]
            # Run blocking embedding API call in thread pool
            batch_embs = await loop.run_in_executor(None, embed_batch, batch)
            embeddings.extend(batch_embs)
            yield f"data: {json.dumps({'type': 'embedding', 'current': min(i + BATCH, total), 'total': total})}\n\n"

        # Save chunks
        for sentence, window, page, emb in zip(all_sentences, all_windows, all_pages, embeddings):
            db.add(Chunk(
                document_id=doc.id,
                content=sentence,
                window_content=window,
                page=page,
                embedding=emb,
            ))

        await db.commit()
        await db.refresh(doc)

        yield f"data: {json.dumps({'type': 'done', 'doc_id': str(doc.id), 'filename': filename, 'chunks': total, 'uploaded_at': doc.uploaded_at.isoformat()})}\n\n"

    except Exception as e:
        yield f"data: {json.dumps({'type': 'error', 'error': str(e)})}\n\n"
    finally:
        os.unlink(tmp_path)


# --- Non-streaming ingestion (kept for backward compat) ---

async def ingest_file(file_bytes: bytes, filename: str, db: AsyncSession, project_id: str | None = None) -> Document:
    suffix = Path(filename).suffix.lower().lstrip(".")
    with tempfile.NamedTemporaryFile(suffix=f".{suffix}", delete=False) as f:
        f.write(file_bytes)
        tmp_path = f.name

    try:
        pages = extract_text(tmp_path, suffix)
        doc = Document(filename=filename, file_type=suffix, project_id=project_id)
        db.add(doc)
        await db.flush()

        all_sentences: list[str] = []
        all_windows: list[str] = []
        all_pages: list[int] = []
        for page_num, page_text in pages:
            for sentence, window in sentence_window_chunks(page_text):
                all_sentences.append(sentence)
                all_windows.append(window)
                all_pages.append(page_num)

        if all_sentences:
            embeddings = embed_texts(all_sentences)
            for sentence, window, page, emb in zip(all_sentences, all_windows, all_pages, embeddings):
                db.add(Chunk(
                    document_id=doc.id,
                    content=sentence,
                    window_content=window,
                    page=page,
                    embedding=emb,
                ))

        await db.commit()
        await db.refresh(doc)
        return doc
    finally:
        os.unlink(tmp_path)
